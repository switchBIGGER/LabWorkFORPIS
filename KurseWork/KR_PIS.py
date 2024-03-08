from tkinter import Tk, Label, Entry, Button, messagebox, ttk, Scrollbar, Text
import psycopg2
from tkcalendar import DateEntry
from datetime import date
import random
import tkinter as tk
import os
from docx import Document

admin_credentials = {
    "admin1": "password1",
    "admin2": "password2"
}

hostname = '127.0.0.1'
username = 'postgres'
password = '1234'
database = 'KR'

class LoginWindow:
    def __init__(self):
        self.root = Tk()
        self.root.title("Окно входа")
        self.init_login()

    def init_login(self):
        self.label = Label(self.root, text="Вы клиент или админ?")
        self.label.grid(row=0, column=0, columnspan=2)

        self.client_button = Button(self.root, text="Клиент", command=self.open_client_window)
        self.client_button.grid(row=1, column=0, padx=5, pady=5)

        self.admin_button = Button(self.root, text="Админ", command=self.open_admin_login)
        self.admin_button.grid(row=1, column=1, padx=5, pady=5)

    def open_client_window(self):
        self.root.destroy()
        client_window = ClientWindow(db2,db)
        client_window.run()

    def open_admin_login(self):
        self.root.destroy()
        admin_login_window = AdminLoginWindow()
        admin_login_window.run()

class AdminLoginWindow:
    def __init__(self):
        self.root = Tk()
        self.root.title("Окно входа для админа")
        self.init_admin()

    def init_admin(self):
        self.login_label = Label(self.root, text="Логин:")
        self.login_label.grid(row=0, column=0, padx=5, pady=5)
        self.login_entry = Entry(self.root)
        self.login_entry.grid(row=0, column=1, padx=5, pady=5)

        self.password_label = Label(self.root, text="Пароль:")
        self.password_label.grid(row=1, column=0, padx=5, pady=5)
        self.password_entry = Entry(self.root, show="*")  
        self.password_entry.grid(row=1, column=1, padx=5, pady=5)

        self.login_button = Button(self.root, text="Войти", command=self.login)
        self.login_button.grid(row=2, column=0, columnspan=2, padx=5, pady=5)

    def login(self):
        login = self.login_entry.get()
        password = self.password_entry.get()

        if login in admin_credentials and admin_credentials[login] == password:
            self.root.destroy()
            # root = tk.Tk()
            AdminWindow(db)
        else:
            messagebox.showerror("Error", "Неверный логин или пароль!")

    def run(self):
        self.root.mainloop()

class AdminWindow(tk.Tk):
    def __init__(self, db):
        tk.Tk.__init__(self)
        self.db = db
        self.init_main()
        self.view_records()
        self.deleted_id = None 

    def init_main(self):
        toolbar = tk.Frame(bg='#d7d8e0', bd=2, master=self)
        toolbar.pack(side=tk.TOP, fill=tk.X)

        btn_open_dialog = tk.Button(toolbar, text='Добавить позицию', command=self.open_dialog, bg='#d7d8e0', bd=0,
                                    compound=tk.TOP)
        btn_open_dialog.pack(side=tk.LEFT)

        btn_edit_dialog = tk.Button(toolbar, text='Редактировать', bg='#d7d8e0', bd=0,
                                    compound=tk.TOP, command=self.open_update_dialog)
        btn_edit_dialog.pack(side=tk.LEFT)

        btn_delete = tk.Button(toolbar, text='Удалить позицию', bg='#d7d8e0', bd=0,
                               compound=tk.TOP, command=self.delete_records)
        btn_delete.pack(side=tk.LEFT)

        self.tree = ttk.Treeview(self, columns=('ID','Surname', 'Name', 'Secondname', 'Money', 'Quantity'), height=15, show='headings')

        self.tree.column('ID', width=120, anchor=tk.CENTER)
        self.tree.column('Surname', width=120, anchor=tk.CENTER)
        self.tree.column('Name', width=100, anchor=tk.CENTER)
        self.tree.column('Secondname', width=100, anchor=tk.CENTER)
        self.tree.column('Money', width=80, anchor=tk.CENTER)
        self.tree.column('Quantity', width=80, anchor=tk.CENTER)


        self.tree.heading('ID', text='ID')
        self.tree.heading('Surname', text='Фамилия')
        self.tree.heading('Name', text='Имя')
        self.tree.heading('Secondname', text='Отчество')
        self.tree.heading('Money', text='Деньги на счету')
        self.tree.heading('Quantity', text='Количество')

        self.tree.pack(side=tk.LEFT)

        scroll = tk.Scrollbar(self, command=self.tree.yview)
        scroll.pack(side=tk.LEFT, fill=tk.Y)
        self.tree.configure(yscrollcommand=scroll.set)

    def records(self, Surname, Name, Secondname, Money, Quantity):
        if self.deleted_id is not None:
            self.db.insert_data_with_id(self.deleted_id, Surname, Name, Secondname, Money, Quantity)
            self.deleted_id = None
        else:
            self.db.insert_data(Surname, Name, Secondname, Money, Quantity)
        self.view_records()

    def update_record(self, Surname, Name, Secondname, money, quantity):
        self.db.c.execute(f'''UPDATE {self.db.name} SET Surname=%s, Name=%s, Secondname=%s, money=%s, quantity=%s WHERE ID=%s''',
                  (Surname, Name, Secondname, money, quantity, self.tree.set(self.tree.selection()[0], '#1')))
        self.db.conn.commit()
        self.view_records()

    def view_records(self):
        self.db.c.execute(f'''SELECT * FROM {self.db.name}''')
        [self.tree.delete(i) for i in self.tree.get_children()]
        [self.tree.insert('', 'end', values=row)
         for row in self.db.c.fetchall()]
        self.tree.get_children()

    def delete_records(self):   
        for selection_item in self.tree.selection():
            deleted_id = self.tree.set(selection_item, '#1')
            self.db.c.execute(f'''DELETE FROM {self.db.name} WHERE id=%s''', (deleted_id,))

        self.db.conn.commit()

        self.view_records()
    def run(self):
        self.root.mainloop()

    def open_dialog(self):
        Child(self)

    def open_update_dialog(self):
        Update(self)


class Child(tk.Toplevel):
    def __init__(self, master):
        super().__init__(master)
        self.init_child()
        self.view = master

    def init_child(self):
        self.title('Добавить пользователя')
        self.geometry('400x360+90+150')
        self.resizable(False, False)

        label_description = tk.Label(self, text='Фамилия:')
        label_description.place(x=50, y=50)
        label_select = tk.Label(self, text='Имя:')
        label_select.place(x=50, y=80)
        label_sum = tk.Label(self, text='Отчество:')
        label_sum.place(x=50, y=110)
        label_select = tk.Label(self, text='Текущий счет:')
        label_select.place(x=50, y=140)
        label_sum = tk.Label(self, text='Количество:')
        label_sum.place(x=50, y=170)

        self.entry_surname = ttk.Entry(self)
        self.entry_surname.place(x=200, y=50)

        self.entry_name = ttk.Entry(self)
        self.entry_name.place(x=200, y=80)

        self.entry_secondname = ttk.Entry(self)
        self.entry_secondname.place(x=200, y=110)

        self.entry_sum = ttk.Entry(self)
        self.entry_sum.place(x=200, y=140)

        self.entry_quantity = ttk.Entry(self)
        self.entry_quantity.place(x=200, y=170)


        btn_cancel = ttk.Button(self, text='Закрыть', command=self.destroy)
        btn_cancel.place(x=300, y=320)

        self.btn_ok = ttk.Button(self, text='Добавить')
        self.btn_ok.place(x=220, y=320)

        self.btn_ok.bind('<Button-1>', lambda event: self.view.records(self.entry_surname.get(),
                                                                       self.entry_name.get(),
                                                                       self.entry_secondname.get(),
                                                                       self.entry_sum.get(),
                                                                       self.entry_quantity.get(),) if len(self.entry_surname.get()) != 0 and len(self.entry_name.get()) != 0 and
                         len(self.entry_secondname.get()) != 0 and len(self.entry_sum.get()) != 0 and
                         len(self.entry_quantity.get()) and
                         self.entry_sum.get().isdigit() and self.entry_quantity.get().isdigit() 
            else messagebox.showerror("Error", "Данные введены неверно!"))

        self.grab_set()
        self.focus_set()


class Update(Child):
    def __init__(self, master):
        super().__init__(master)
        self.init_edit()
        self.view = master

    def init_edit(self):
        self.title('Редактировать позицию')
        btn_edit = ttk.Button(self, text='Редактировать')
        btn_edit.place(x=205, y=320)
        btn_edit.bind('<Button-1>', lambda event: self.view.update_record(self.entry_surname.get(),
                                                                          self.entry_name.get(),
                                                                          self.entry_secondname.get(),
                                                                          self.entry_sum.get(),
                                                                          self.entry_quantity.get(),) if len(self.entry_name.get()) != 0 and len(self.entry_surname.get()) != 0 and
                      len(self.entry_secondname.get()) != 0 and len(self.entry_sum.get()) != 0 and
                      len(self.entry_quantity.get()) != 0 and
                      self.entry_sum.get().isdigit() and self.entry_quantity.get().isdigit() 
            else messagebox.showerror("Error", "Данные введены неверно!"))

        self.btn_ok.destroy()

class ClientWindow:
    def __init__(self, db1, db2):
        self.root = Tk()
        self.root.title("АО Лошаков")
        self.db1 = db1
        self.db2 = db2
        self.init_client()
        self.view_records()

    def init_client(self):
        self.combobox_left = ttk.Combobox(self.root)
        self.combobox_left.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.combobox_left['values'] = ("Продажа", "Продажа за наличный расчет", "Возврат товара от покупателя", "Отчет о продажах")
        self.combobox_left.set("Выбери вид продажи")

        self.client_label = Label(self.root, text="Клиент:")
        self.client_label.grid(row=0, column=1, padx=5, pady=5, sticky="e")

        self.combobox = ttk.Combobox(self.root)
        self.combobox.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.combobox['values'] = clients
        self.combobox.set(clients[0])

        self.reason_label = Label(self.root, text="Основание:")
        self.reason_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")

        self.reason_entry = Entry(self.root)
        self.reason_entry.grid(row=1, column=1, padx=5, pady=5, columnspan=2, sticky="we")

        random_text = "Товарный чек:\n"
        random_text += "".join(random.choices("abcdefghijklmnopqrstuvwxyz1234567890", k=20))

        self.receipt_label = Label(self.root, text=random_text, wraplength=150)
        self.receipt_label.grid(row=1, column=3, padx=5, pady=5, sticky="e")

        self.date_label = Label(self.root, text="Дата:")
        self.date_label.grid(row=1, column=4, padx=5, pady=5, sticky="e")

        self.date_entry = DateEntry(self.root, width=12, date_pattern="dd.mm.yyyy")
        self.date_entry.grid(row=1, column=5, padx=5, pady=5, sticky="w")

        self.quantity_label = Label(self.root, text="Количество:")
        self.quantity_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")

        self.quantity_entry = Entry(self.root)
        self.quantity_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")

        self.treeview = ttk.Treeview(self.root)
        self.treeview["columns"] = ("id", "name", "quantity", "unit", "price", "sum")

        self.treeview.heading("id", text="ID")
        self.treeview.heading("name", text="Наименование")
        self.treeview.heading("quantity", text="Количество")
        self.treeview.heading("unit", text="Ед. изм.")
        self.treeview.heading("price", text="Цена")
        self.treeview.heading("sum", text="Сумма")

        self.treeview["show"] = "headings"

        self.treeview.grid(row=3, column=0, padx=5, pady=5, columnspan=7)

        self.scrollbar = Scrollbar(self.root, orient="vertical", command=self.treeview.yview)
        self.scrollbar.grid(row=3, column=8, sticky="ns")

        self.treeview.configure(yscrollcommand=self.scrollbar.set)

        self.ok_button = Button(self.root, text="OK", command=self.ok_button_clicked)
        self.ok_button.grid(row=4, column=0, padx=5, pady=5, columnspan=7)

    def view_records(self):
        self.db1.c.execute(f'''SELECT * FROM {self.db1.name}''')
        [self.treeview.delete(i) for i in self.treeview.get_children()]
        [self.treeview.insert('', 'end', values=row) for row in self.db1.c.fetchall()]
        self.treeview.get_children()

    def ok_button_clicked(self):
        selected_option = self.combobox_left.get()
        if selected_option == "Продажа":
            self.create_invoice()
        elif selected_option == "Продажа за наличный расчет":
            self.create_cash_sale()
        elif selected_option == "Возврат товара от покупателя":
            self.create_return()
        elif selected_option == "Отчет о продажах":
            self.generate_sales_report()
        else:
            messagebox.showerror("Ошибка", "Выберите правильный тип продажи.")
    def create_invoice(self):
        sale_type = self.combobox_left.get()
        client_name = self.combobox.get()
        product = self.treeview.item(self.treeview.focus())["values"]
        quantity = self.quantity_entry.get()
        price = product[4]

        document_name = f"Счет_на_оплату_{client_name}_{date.today()}.txt"
        with open(document_name, 'w', encoding='utf-8') as file:
            file.write(f"Тип продажи: {sale_type}\n")
            file.write(f"Имя клиента: {client_name}\n")
            file.write(f"Дата: {date.today()}\n")
            file.write(f"Товар: {product[1]}\n")
            file.write(f"Количество: {quantity}\n")
            file.write(f"Цена: {price}\n")
            file.write(f"Итого: {int(quantity) * price}\n")

        messagebox.showinfo("Успех", "Документ успешно создан.")

    def create_cash_sale(self):
        sale_type = self.combobox_left.get()
        client_name = self.combobox.get()
        product = self.treeview.item(self.treeview.focus())["values"]
        quantity = self.quantity_entry.get()
        price = product[4]
        date = self.date_entry.get()

        document_name = f"Товарный_чек_{client_name}_{date}.txt"
        receipt_number = random.randint(100000, 999999)
        receipt_text = f"Товарный чек: {receipt_number}"

        with open(document_name, 'w', encoding='utf-8') as file:
            file.write(f"Тип продажи: {sale_type}\n")
            file.write(f"Имя клиента: {client_name}\n")
            file.write(f"Основание: {receipt_text}\n")
            file.write(f"Дата: {date}\n")
            file.write(f"Товар: {product[1]}\n")
            file.write(f"Количество: {quantity}\n")
            file.write(f"Цена: {price}\n")
            file.write(f"Итого: {int(quantity) * price}\n")

        messagebox.showinfo("Успех", "Документ успешно создан.")

    def create_return(self):
        sale_type = self.combobox_left.get()
        client_name = self.combobox.get()
        return_date = date.today()
        return_reason = self.reason_entry.get()
        product = self.treeview.item(self.treeview.focus())["values"]
        price = product[4]
        return_quantity = self.quantity_entry.get()

        document_name = f"Возврат_товара_{client_name}_{return_date}.txt"
        sale_document = self.receipt_label.cget("text")

        with open(document_name, 'w', encoding='utf-8') as file:
            file.write(f"Тип продажи: {sale_type}\n")
            file.write(f"Имя клиента: {client_name}\n")
            file.write(f"Дата: {return_date}\n")
            file.write(f"Номер документа продажи: {sale_document}\n")
            file.write(f"Причина возврата: {return_reason}\n")
            file.write(f"Товар: {product[1]}\n")
            file.write(f"Количество: {return_quantity}\n")
            file.write(f"Цена: {price}\n")
            file.write(f"Итого: {int(return_quantity) * price}\n")

        messagebox.showinfo("Успех", "Документ успешно создан.")

    def generate_sales_report(self):
        sale_type = self.combobox.get()
        start_date = self.date_entry.get()

        report_name = f"Отчет_о_продажах_{sale_type}_{start_date}.docx"
        document = Document()
        document.add_heading("Отчет о продажах", level=1)

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Вид продажи"
        table.cell(0, 1).text = "Дата продажи"
        table.cell(0, 2).text = "Сумма"

        for filename in os.listdir():
            if filename.startswith("Счет_на_оплату") or filename.startswith("Товарный_чек") or filename.startswith("Возврат_товара"):
                with open(filename, 'r', encoding='utf-8') as sales_file:
                    sales_data = sales_file.readlines()

                    sale_type_file = ""
                    sale_date = ""
                    sale_amount = ""

                    for line in sales_data:
                        if line.startswith("Тип продажи"):
                            sale_type_file = line.split(':')[1].strip()
                        elif line.startswith("Дата"):
                            sale_date = line.split(':')[1].strip()
                        elif line.startswith("Итого"):
                            sale_amount = line.split(':')[1].strip()

                    row_cells = table.add_row().cells
                    row_cells[0].text = sale_type_file
                    row_cells[1].text = sale_date
                    row_cells[2].text = sale_amount

        document.save(report_name)

        messagebox.showinfo("Успех", "Отчет о продажах успешно сгенерирован.")
    def run(self):
        self.root.mainloop()


class DB_users():
    def __init__(self, name):
        self.name = name
        self.conn = psycopg2.connect(host=hostname,
        user=username,
        password=password,
        database=database)
        self.c = self.conn.cursor()
        self.c.execute(
            f'''CREATE TABLE IF NOT EXISTS {self.name} (id serial primary key, Surname text, Name text,
              Secondname text, money integer, quantity integer)''')
        self.conn.commit()

    def insert_data(self, Surname, Name, Secondname, money, quantity):
        self.c.execute(f'''INSERT INTO {self.name}(Surname, Name, Secondname, money, quantity) VALUES (%s, %s, %s, %s, %s)''',
                       (Surname, Name, Secondname, money, quantity))
        self.conn.commit()
    def get_all_names(self):
        self.c.execute("SELECT Surname, Name, Secondname FROM admins")
        names = [f"{row[0]} {row[1][0]}.{row[2][0]}." for row in self.c.fetchall()]
        return names
class DB_products():
    def __init__(self, name):
        self.name = name
        self.conn = psycopg2.connect(host=hostname,
        user=username,
        password=password,
        database=database)
        self.c = self.conn.cursor()
        self.c.execute(
            f'''CREATE TABLE IF NOT EXISTS {self.name} (id serial primary key, name text, quantity integer, ed text, price integer, summ integer)''')
        self.conn.commit()

    def insert_data(self, name, quantity, ed, price):
        nds = 0.2 
        summ = price + (price * nds)

        self.c.execute(f'''INSERT INTO {self.name}(name, quantity, ed, price, summ) VALUES (%s, %s, %s, %s, %s)''',
                       (name, quantity, ed, price, summ))
        self.conn.commit()


if __name__ == '__main__':
    login_window = LoginWindow()
    db = DB_users('admins')
    db2 = DB_products('product_1')
    clients = db.get_all_names()
    login_window.root.mainloop()

